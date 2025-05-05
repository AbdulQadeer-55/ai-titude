import os
import json
import docx
import tempfile
import requests
import logging
import hashlib
import time
import re
from django.http import JsonResponse, HttpResponse
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.core.cache import cache
from rest_framework.decorators import api_view
from rest_framework import status
from google.cloud import documentai_v1 as documentai
from google.cloud import texttospeech
from google.api_core.client_options import ClientOptions
import google.generativeai as genai
from google.oauth2 import service_account
from dotenv import load_dotenv
from requests.exceptions import RequestException
from pydub import AudioSegment
from django.views.decorators.csrf import csrf_exempt

load_dotenv()

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

PROJECT_ID = os.getenv('PROJECT_ID', '585624190112')
LOCATION = os.getenv('LOCATION', 'us')
PROCESSOR_ID = os.getenv('PROCESSOR_ID', '867de035aee8db09')
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', '')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')
LOUDLY_API_KEY = os.getenv('LOUDLY_API_KEY', 'yfZX66sUgEGqQLrHpI-ysbcTNrMozGnXtqPlyD-6NXQ')
SECRET_KEY = os.getenv('SECRET_KEY', '')
GPT4O_MINI_TTS_API_KEY = os.getenv('OPENAI_API_KEY', '')
genai.configure(api_key=GOOGLE_API_KEY)

MAX_FILE_SIZE_ONLINE = 20 * 1024 * 1024
MAX_FILE_SIZE_BATCH = 1 * 1024 * 1024 * 1024
MAX_IMAGE_RESOLUTION = 40 * 1000000
MAX_FILES_PER_BATCH = 5000
MAX_PAGES_ONLINE_DEFAULT = 15
MAX_PAGES_IMAGELESS = 30
MAX_TEXT_LENGTH = 10000

EMOTION_MAPPING = {
    "neutral": "neutral",
    "sympathetic": "sympathetic",
    "sincere": "sincere",
    "calm": "calm",
    "serene": "serene",
    "sadness": "sad",
    "happiness": "happy",
    "fear": "fearful",
    "horror": "horrified",
    "surprise": "surprised",
    "anger": "angry",
    "rage": "enraged",
    "love": "loving",
    "excitement": "excited",
    "anxiety": "anxious",
    "disgust": "disgusted"
}

def load_urdu_dictionary(file_path=os.path.join(settings.BASE_DIR, 'api', 'urdu_words.txt')):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f if line.strip())
    except FileNotFoundError:
        logger.error(f"Could not find Urdu dictionary at {file_path}")
        return set()

urdu_dictionary = load_urdu_dictionary()

def create_error_response(code, message, details=None):
    response = {'error': {'code': code, 'message': message}}
    if details:
        response['error']['details'] = details
    logger.error(f"Error {code}: {message} - Details: {details}")
    return JsonResponse(response, status=code)


def create_error_response(code, message, details=None):
    response = {'error': {'code': code, 'message': message}}
    if details:
        response['error']['details'] = details
    logger.error(f"Error {code}: {message} - Details: {details}")
    return JsonResponse(response, status=code)

def process_document(file_path, mime_type, is_batch=False):
    try:
        credentials_path = "/home/abdulqadeer/Downloads/ai-titude/credentials.json"
        if not os.path.exists(credentials_path):
            raise FileNotFoundError("Google Cloud credentials file not found")
        credentials = service_account.Credentials.from_service_account_file(credentials_path)
        client_options = ClientOptions(api_endpoint=f"{LOCATION}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options, credentials=credentials)
        name = client.processor_path(PROJECT_ID, LOCATION, PROCESSOR_ID)
        
        with open(file_path, "rb") as f:
            file_content = f.read()

        file_size = len(file_content)
        max_file_size = MAX_FILE_SIZE_BATCH if is_batch else MAX_FILE_SIZE_ONLINE
        if file_size > max_file_size:
            limit_mb = max_file_size // (1024 * 1024)
            raise ValueError(f"File size exceeds the {'batch' if is_batch else 'online'} processing limit of {limit_mb} MB.")

        raw_document = documentai.RawDocument(content=file_content, mime_type=mime_type)
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)
        result = client.process_document(request=request)
        document = result.document

        page_count = len(document.pages)
        max_pages = MAX_PAGES_IMAGELESS if 'imageless_mode' in request else MAX_PAGES_ONLINE_DEFAULT
        if page_count > max_pages:
            raise ValueError(f"Document exceeds the page limit of {max_pages} pages for {'imageless' if 'imageless_mode' in request else 'standard'} mode.")
        return document
    except FileNotFoundError as e:
        logger.error(f"Document AI error: {str(e)}")
        raise
    except ValueError as ve:
        logger.error(f"Document AI validation error: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Failed to process document with Document AI: {str(e)}")
        raise

def format_extracted_text(document):
    return document.text

def extract_text_from_file(file_path, file_extension):
    try:
        mime_type_map = {
            'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
            'bmp': 'image/bmp', 'tiff': 'image/tiff', 'tif': 'image/tiff',
            'gif': 'image/gif', 'pdf': 'application/pdf',
        }
        if file_extension in mime_type_map:
            mime_type = mime_type_map[file_extension]
            document = process_document(file_path, mime_type)
            if document:
                return format_extracted_text(document)
        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError("Unsupported file type. Supported types are: jpg, png, pdf, docx, txt.")
    except ValueError as ve:
        logger.error(f"Text extraction error: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Error extracting text from file: {str(e)}")
        raise

def get_gemini_response(prompt, document_text):
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        content = [document_text, prompt]
        response = model.generate_content(content)
        if not response.text:
            raise ValueError("Gemini API returned an empty response.")
        return response.text
    except ValueError as ve:
        logger.error(f"Gemini error: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Error with Gemini API: {str(e)}")
        raise

def filter_unethical_text_with_prompt(text):
    prompt = f"""
    Analyze the following text and remove any vulgar, abusive, or unethical words or phrases in Urdu or any other language.
    Return only the cleaned text without any explanations or markers, preserving the original meaning as much as possible.
    If no unethical content is found, return the text as is.
    Text: "{text}"
    """
    try:
        cleaned_text = get_gemini_response(prompt, text)
        return cleaned_text.strip() if cleaned_text else text
    except Exception as e:
        logger.error(f"Failed to filter unethical text: {str(e)}")
        raise

def detect_emotions_with_ai(text):
    emotion_list = list(EMOTION_MAPPING.keys())
    prompt = f"""
    Analyze the following text and identify the single most prominent emotion from this list: {', '.join(emotion_list)}.
    The text may be in Urdu, English, or a mix of both or any other language. Provide only the most dominant emotion as a single word or phrase.
    Text: "{text}"
    """
    try:
        response = get_gemini_response(prompt, text)
        return response.strip() if response else "No clear emotion detected"
    except Exception as e:
        logger.error(f"Failed to detect emotion: {str(e)}")
        raise

def detect_gender_with_ai(text):
    prompt = f"""
    Analyze the following Urdu text and determine the gender of the subject (male, female, or unknown).
    Look for gender-specific indicators such as pronouns (e.g., "وہ" with context), verb conjugations (e.g., "گیا" for male, "گئی" for female),
    or nouns (e.g., "لڑکا" for male, "لڑکی" for female). Return only the detected gender as "male", "female", or "unknown".
    Text: "{text}"
    """
    try:
        response = get_gemini_response(prompt, text)
        gender = response.strip().lower()
        if gender not in ["male", "female", "unknown"]:
            return "unknown"
        return gender
    except Exception as e:
        logger.error(f"Failed to detect gender: {str(e)}")
        raise

def get_available_voices():
    try:
        credentials_path = "/home/abdulqadeer/Downloads/ai-titude/credentials.json"
        if not os.path.exists(credentials_path):
            raise FileNotFoundError("Google Cloud credentials file not found")
        credentials = service_account.Credentials.from_service_account_file(credentials_path)
        client = texttospeech.TextToSpeechClient(credentials=credentials)
        voices = client.list_voices().voices
        language_voices = {}
        for voice in voices:
            for language_code in voice.language_codes:
                if language_code not in language_voices:
                    language_voices[language_code] = []
                language_voices[language_code].append({
                    'name': voice.name,
                    'gender': texttospeech.SsmlVoiceGender(voice.ssml_gender).name,
                    'language_code': language_code
                })
        sorted_languages = sorted(language_voices.keys())
        result = []
        for lang in sorted_languages:
            result.append({
                'language_code': lang,
                'voices': sorted(language_voices[lang], key=lambda x: (x['gender'], x['name']))
            })
        gpt4o_voices = [
            {
                'language_code': 'ur-PK',
                'voices': [
                    {'name': 'alloy', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'ash', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'ballad', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'coral', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'echo', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'fable', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'onyx', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'nova', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'sage', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'},
                    {'name': 'shimmer', 'gender': 'NEUTRAL', 'language_code': 'ur-PK'}      
                ]
            }
        ]
        result.extend(gpt4o_voices)
        return result
    except FileNotFoundError as e:
        logger.error(f"Voice fetch error: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"Error fetching available voices: {str(e)}")
        raise

@csrf_exempt
def generate_music_with_prompt(request):
    """
    Generate music based on a user-provided prompt and duration using the Loudly API.
    """
    if request.method != "POST":
        return create_error_response(405, "Only POST requests allowed")

    try:
        # Parse request body
        body_unicode = request.body.decode("utf-8")
        body = json.loads(body_unicode)
        prompt = body.get("prompt")
        duration = body.get("duration", 90)  # Default to 90 seconds
        logger.debug(f"Incoming request: prompt={prompt}, duration={duration}")

        # Validate inputs
        if prompt is None:
            return create_error_response(400, "Prompt is missing", "The 'prompt' field is required in the request body")
        if not isinstance(prompt, str):
            return create_error_response(400, "Invalid prompt type", "Prompt must be a string")
        prompt = prompt.strip()
        if not prompt:
            return create_error_response(400, "Prompt is empty", "Please provide a valid music prompt")
        if len(prompt) < 5:
            return create_error_response(400, "Prompt is too short", "Minimum length is 5 characters")
        if not isinstance(duration, int) or not (30 <= duration <= 420):
            return create_error_response(400, "Invalid duration", "Duration must be an integer between 30 and 420 seconds")

        # Check Loudly API key
        if not LOUDLY_API_KEY:
            return create_error_response(500, "Loudly API key is not configured")

        # Common headers
        headers = {
            "API-KEY": LOUDLY_API_KEY,
            "Content-Type": "application/json",
            "Accept": "application/json",
            "User-Agent": "AI-titude/1.0",
            "Cache-Control": "no-cache"
        }

        # Try multiple payload formats
        payload_variations = [
            {"prompt": prompt, "duration": duration},
            {"Prompt": prompt, "duration": duration},
            {"data": {"prompt": prompt, "duration": duration}}
        ]

        # Try multiple endpoints
        endpoints = [
            "https://soundtracks-dev.loudly.com/api/ai/prompt/songs",
            "https://soundtracks.loudly.com/b2b/ai/prompt/songs"
        ]

        response = None
        for endpoint in endpoints:
            for payload in payload_variations:
                logger.debug(f"Attempting Loudly API: URL={endpoint}, headers={headers}, payload={payload}")
                try:
                    response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
                    logger.debug(f"Loudly API response: status={response.status_code}, text={response.text}")
                    
                    if response.status_code == 200:
                        try:
                            response_data = response.json()
                            if not response_data.get("music_file_path"):
                                logger.error("Loudly API response missing music_file_path")
                                return create_error_response(500, "Invalid response from Loudly API", "Missing music_file_path")
                            logger.info(f"Music generated successfully: {response_data}")
                            return JsonResponse(response_data, status=200)
                        except ValueError:
                            logger.error("Invalid JSON response from Loudly API")
                            return create_error_response(500, "Invalid JSON response from Loudly API")
                    
                    elif response.status_code != 400:
                        # If not a 400 error, stop trying this endpoint
                        break
                    
                except requests.RequestException as e:
                    logger.error(f"Network error calling Loudly API at {endpoint}: {str(e)}")
                    continue

            if response and response.status_code != 200:
                try:
                    error_data = response.json()
                    logger.error(f"Loudly API error at {endpoint}: {error_data}")
                    return create_error_response(response.status_code, "Failed to generate music", error_data.get("error", response.text))
                except ValueError:
                    logger.error(f"Loudly API error at {endpoint}: status={response.status_code}, text={response.text}")
                    return create_error_response(response.status_code, "Failed to generate music", response.text)

        # If all JSON attempts fail, try form data as a fallback
        form_headers = {
            "API-KEY": LOUDLY_API_KEY,
            "Accept": "application/json",
            "User-Agent": "AI-titude/1.0",
            "Cache-Control": "no-cache"
        }
        form_data = {"prompt": prompt, "duration": str(duration)}
        logger.debug(f"Attempting Loudly API with form data: URL={endpoints[0]}, headers={form_headers}, form_data={form_data}")
        try:
            response = requests.post(endpoints[0], headers=form_headers, data=form_data, timeout=30)
            logger.debug(f"Loudly API form data response: status={response.status_code}, text={response.text}")
            
            if response.status_code == 200:
                try:
                    response_data = response.json()
                    if not response_data.get("music_file_path"):
                        logger.error("Loudly API response missing music_file_path")
                        return create_error_response(500, "Invalid response from Loudly API", "Missing music_file_path")
                    logger.info(f"Music generated successfully: {response_data}")
                    return JsonResponse(response_data, status=200)
                except ValueError:
                    logger.error("Invalid JSON response from Loudly API")
                    return create_error_response(500, "Invalid JSON response from Loudly API")
            
            try:
                error_data = response.json()
                logger.error(f"Loudly API form data error: {error_data}")
                return create_error_response(response.status_code, "Failed to generate music", error_data.get("error", response.text))
            except ValueError:
                logger.error(f"Loudly API form data error: status={response.status_code}, text={response.text}")
                return create_error_response(response.status_code, "Failed to generate music", response.text)
        
        except requests.RequestException as e:
            logger.error(f"Network error calling Loudly API with form data: {str(e)}")
            return create_error_response(500, "Network error", str(e))

        # If all attempts fail
        return create_error_response(400, "Failed to generate music", "All attempts to contact Loudly API failed. Please check API configuration.")

    except json.JSONDecodeError:
        logger.error("Invalid JSON in request body")
        return create_error_response(400, "Invalid JSON format in request body")
    except requests.RequestException as e:
        logger.error(f"Network error calling Loudly API: {str(e)}")
        return create_error_response(500, "Network error", str(e))
    except Exception as e:
        logger.error(f"Unexpected error in generate_music_with_prompt: {str(e)}", exc_info=True)
        return create_error_response(500, "An unexpected error occurred", str(e))




def text_to_speech(text, voice_settings_list, base_file_name, detected_gender, tts_provider="google"):
    logger.debug(f"Starting text_to_speech: provider={tts_provider}, text_length={len(text)}, voice_settings={voice_settings_list}")
    try:
        if len(text) > MAX_TEXT_LENGTH:
            raise ValueError(f"Text exceeds the maximum length of {MAX_TEXT_LENGTH} characters for audio generation.")
        
        voice_settings = voice_settings_list[0]
        
        if tts_provider == "gpt4o_mini":
            logger.debug("Using GPT-4o mini TTS")
            required_fields = ['voice_name']
            if not all(key in voice_settings for key in required_fields):
                raise ValueError(f"Invalid voice settings: Missing required fields: {required_fields}")
            
            valid_voices = ['alloy', 'ash', 'ballad', 'coral', 'echo', 'fable', 'onyx', 'nova', 'sage', 'shimmer']
            if voice_settings['voice_name'] not in valid_voices:
                raise ValueError(f"Invalid voice: {voice_settings['voice_name']}. Supported voices: {valid_voices}")
            
            valid_emotions = list(EMOTION_MAPPING.keys())
            instructions = voice_settings.get('instructions', '').strip()
            
            if not instructions:
                raise ValueError("Instructions are required for GPT-4o mini TTS.")
            emotion_match = re.search(r'Speak in a (\w+) tone with (\d+)% intensity', instructions)
            secondary_emotion_match = re.search(r'blended with (\w+) at (\d+)% intensity', instructions)
            tone_match_standard = re.search(r',\s*(\w+)\s*tone,\s*', instructions)
            style_match = re.search(r'(\w+(?:-\w+)*) style', instructions)
            pacing_match = re.search(r'pacing at (\d+)%', instructions)
            pause_frequency_match = re.search(r'pause frequency set to (\w+)', instructions)
            emphasis_match = re.search(r'Emphasize the following words: ([\w\s,آ-ی]+)\.', instructions)

            if not emotion_match:
                alt_emotion_tone_match = re.search(r'Speak with a ([\w\s,]+) tone', instructions)
                smile_match = re.search(r'and a slight smile', instructions)

                if alt_emotion_tone_match:
                    tone_description = alt_emotion_tone_match.group(1).lower()
                    tone_mapping = {
                        "bright, cheerful": "excited",
                        "cheerful": "excited",
                        "bright": "excited",
                        "calm": "calm",
                        "soothing": "soothing",
                    }
                    emotion_mapping_alt = {
                        "bright, cheerful": "happiness",
                        "cheerful": "happiness",
                        "bright": "happiness",
                        "calm": "calm",
                        "soothing": "serene",
                    }
                    tone = tone_mapping.get(tone_description, "excited")
                    emotion = emotion_mapping_alt.get(tone_description, "happiness")
                    if smile_match:
                        emotion = "happiness"
                    emotion_intensity = 70
                    secondary_emotion = "none"
                    secondary_emotion_intensity = 0
                    style = "conversational"
                    pacing = 100
                    pause_frequency = "medium"
                    emphasis_words = None
                else:
                    raise ValueError("Invalid instructions format. Expected format: 'Speak in a {emotion} tone with {intensity}% intensity...' or 'Speak with a {tone} tone and a slight smile'")
            else:
                emotion = emotion_match.group(1).lower() if emotion_match else None
                emotion_intensity = int(emotion_match.group(2)) if emotion_match else 70
                secondary_emotion = secondary_emotion_match.group(1).lower() if secondary_emotion_match else 'none'
                secondary_emotion_intensity = int(secondary_emotion_match.group(2)) if secondary_emotion_match else 0
                tone = tone_match_standard.group(1).lower() if tone_match_standard else 'empathetic'
                style = style_match.group(1).lower() if style_match else 'conversational'
                pacing = int(pacing_match.group(1)) if pacing_match else 100
                pause_frequency = pause_frequency_match.group(1).lower() if pause_frequency_match else 'medium'
                emphasis_words = emphasis_match.group(1).strip() if emphasis_match else None

            if emotion not in valid_emotions:
                raise ValueError(f"Invalid base emotion: {emotion}. Supported emotions: {valid_emotions}")
            if secondary_emotion != 'none' and secondary_emotion not in valid_emotions:
                raise ValueError(f"Invalid secondary emotion: {secondary_emotion}. Supported emotions: {valid_emotions}")

            if not (0 <= emotion_intensity <= 100):
                raise ValueError(f"Base emotion intensity must be between 0 and 100, got {emotion_intensity}")
            if not (0 <= secondary_emotion_intensity <= 100):
                raise ValueError(f"Secondary emotion intensity must be between 0 and 100, got {secondary_emotion_intensity}")

            valid_tones = ['empathetic', 'solution-focused', 'gentle', 'authoritative', 'warm',
                           'soothing', 'excited', 'noble', 'chaotic', 'calm']
            if tone not in valid_tones:
                raise ValueError(f"Invalid tone: {tone}. Supported tones: {valid_tones}")

            valid_styles = ['conversational', 'professional', 'dramatic', 'monotone', 'narrative', 'poetic',
                            'motivational', 'whispered', 'sarcastic', 'childlike', 'commanding', 'meditative',
                            'sports-coach', 'bedtime-story', 'medieval-knight', 'mad-scientist', 'patient-teacher',
                            'auctioneer', 'old-timey', 'chill-surfer']
            if style not in valid_styles:
                raise ValueError(f"Invalid style: {style}. Supported styles: {valid_styles}")

            if not (50 <= pacing <= 200):
                raise ValueError(f"Pacing must be between 50% and 200%, got {pacing}%")

            valid_pause_frequencies = ['low', 'medium', 'high']
            if pause_frequency not in valid_pause_frequencies:
                raise ValueError(f"Invalid pause frequency: {pause_frequency}. Supported values: {valid_pause_frequencies}")

            if emphasis_words:
                words_to_emphasize = [word.strip() for word in emphasis_words.split(',')]
                for word in words_to_emphasize:
                    text = text.replace(word, f'<emphasis level="strong">{word}</emphasis>')

            speed = pacing / 100.0
            if speed < 0.5:
                speed = 0.5
            elif speed > 2.0:
                speed = 2.0

            plain_text = text.replace('.', '. ')
            print(f"Plain Text Input: {plain_text}")

            headers = {
                "Authorization": f"Bearer {GPT4O_MINI_TTS_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "gpt-4o-mini-tts",
                "input": plain_text,
                "voice": voice_settings['voice_name'],
                "response_format": "mp3",
                "speed": speed
            }
            if 'instructions' in voice_settings:
                payload['instructions'] = voice_settings['instructions']
            
            logger.debug(f"GPT-4o mini TTS payload: {payload}")
            
            response = requests.post(
                "https://api.openai.com/v1/audio/speech",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code != 200:
                logger.error(f"GPT-4o mini TTS API error: {response.status_code} - {response.text}")
                raise Exception(f"GPT-4o mini TTS API error: {response.text}")
            
            final_file_name = f"{base_file_name.rsplit('.', 1)[0]}.mp3"
            final_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
            with open(final_temp.name, "wb") as out:
                out.write(response.content)
            logger.debug(f"Saved GPT-4o mini audio to {final_temp.name}")
            
            return final_temp.name, final_file_name, None
        
        else:
            logger.debug("Using Google Cloud TTS")
            required_fields = ['language_code', 'voice_name', 'gender']
            if not all(key in voice_settings for key in required_fields):
                raise ValueError(f"Invalid voice settings: Missing required fields: {required_fields}")
            
            if detected_gender != "unknown" and detected_gender != voice_settings['gender'].lower():
                raise ValueError(f"Gender mismatch: Text is detected as {detected_gender}, but selected voice is {voice_settings['gender'].lower()}.")
            
            credentials_path = "/home/abdulqadeer/Downloads/ai-titude/credentials.json"
            if not os.path.exists(credentials_path):
                raise FileNotFoundError("Google Cloud credentials file not found")
            credentials = service_account.Credentials.from_service_account_file(credentials_path)
            client = texttospeech.TextToSpeechClient(credentials=credentials)
            MAX_BYTES = 5000
            text = text.replace('\n', ' ').strip()
            temp_files = []
            combined = AudioSegment.empty()
            
            voice = texttospeech.VoiceSelectionParams(
                language_code=voice_settings['language_code'],
                name=voice_settings['voice_name'],
                ssml_gender=texttospeech.SsmlVoiceGender[voice_settings['gender'].upper()]
            )
            audio_config = texttospeech.AudioConfig(
                audio_encoding=texttospeech.AudioEncoding.MP3,
                speaking_rate=float(voice_settings.get('speaking_rate', 1.0)),
                pitch=float(voice_settings.get('pitch', 0.0)),
                volume_gain_db=float(voice_settings.get('volume_gain_db', 0.0)),
                effects_profile_id=voice_settings.get('audio_effects', [])
            )
            
            chunks = [text[i:i+MAX_BYTES] for i in range(0, len(text), MAX_BYTES)]
            for i, chunk in enumerate(chunks):
                synthesis_input = texttospeech.SynthesisInput(text=chunk)
                response = client.synthesize_speech(input=synthesis_input, voice=voice, audio_config=audio_config)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_chunk{i}.mp3")
                with open(temp_file.name, "wb") as out:
                    out.write(response.audio_content)
                temp_files.append(temp_file.name)
            
            for temp_file in temp_files:
                audio_segment = AudioSegment.from_mp3(temp_file)
                combined += audio_segment
            
            final_file_name = f"{base_file_name.rsplit('.', 1)[0]}.mp3"
            final_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
            combined.export(final_temp.name, format="mp3")
            for temp_file in temp_files:
                os.unlink(temp_file)
            logger.debug(f"Saved Google Cloud audio to {final_temp.name}")
            
            return final_temp.name, final_file_name, None
    
    except FileNotFoundError as e:
        logger.error(f"TTS error: {str(e)}")
        raise
    except ValueError as ve:
        logger.error(f"TTS validation error: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Error generating audio: {str(e)}")
        raise

@api_view(['POST'])
def analyze_files(request):
    logger.debug("Received analyze_files request")
    try:
        files = request.FILES.getlist('files')
        if not files:
            return create_error_response(400, "No files uploaded.", "Please upload at least one file to analyze.")
        if len(files) > MAX_FILES_PER_BATCH:
            return create_error_response(400, f"Too many files uploaded.", f"Maximum of {MAX_FILES_PER_BATCH} files allowed.")
        
        fs = FileSystemStorage()
        combined_text = []
        for file in files:
            file_extension = file.name.split('.')[-1].lower()
            if file.size > MAX_FILE_SIZE_ONLINE:
                return create_error_response(400, f"File '{file.name}' is too large.", f"Maximum file size is 20 MB.")
            
            filename = fs.save(file.name, file)
            file_path = fs.path(filename)
            try:
                document_text = extract_text_from_file(file_path, file_extension)
                if not document_text:
                    continue
                prompt = """
                Extract content in Urdu only, including all diacritic marks such as zair, zabar, pesh, and all other diacritic marks.
                Output only the extracted Urdu content without explanations.
                """
                extracted_text = get_gemini_response(prompt, document_text)
                if extracted_text:
                    cleaned_text = filter_unethical_text_with_prompt(extracted_text)
                    if cleaned_text:
                        combined_text.append(cleaned_text)
            finally:
                if os.path.exists(file_path):
                    os.remove(file_path)
        
        if not combined_text:
            return create_error_response(400, "No valid text extracted.")
        
        extracted_text = "\n\n".join(combined_text)
        detected_emotion = detect_emotions_with_ai(extracted_text)
        detected_gender = detect_gender_with_ai(extracted_text)
        logger.debug(f"Analyze files response: text_length={len(extracted_text)}, emotion={detected_emotion}, gender={detected_gender}")
        return JsonResponse({
            'extracted_text': extracted_text,
            'detected_emotion': detected_emotion,
            'detected_gender': detected_gender
        })
    except Exception as e:
        logger.error(f"Analyze files error: {str(e)}")
        return create_error_response(500, "An unexpected error occurred.", str(e))

@api_view(['POST'])
def generate_audio(request):
    logger.debug(f"Received generate_audio request: {request.body}")
    try:
        data = json.loads(request.body)
        text = data.get('text', '')
        voice_settings_list = data.get('voice_settings_list', [])
        detected_gender = data.get('detected_gender', 'unknown')
        tts_provider = data.get('tts_provider', 'google')
        
        if not text:
            return create_error_response(400, "No text provided.")
        if not voice_settings_list or not isinstance(voice_settings_list, list):
            return create_error_response(400, "Invalid voice settings.")
        
        base_file_name = "generated_audio"
        audio_file, audio_file_name, warning = text_to_speech(text, voice_settings_list, base_file_name, detected_gender, tts_provider)
        with open(audio_file, 'rb') as f:
            audio_bytes = f.read()
        os.unlink(audio_file)
        
        response = HttpResponse(audio_bytes, content_type='audio/mp3')
        response['Content-Disposition'] = f'attachment; filename="{audio_file_name}"'
        logger.debug(f"Generated audio file: {audio_file_name}")
        return response
    except ValueError as ve:
        logger.error(f"Generate audio validation error: {str(ve)}")
        return create_error_response(400, str(ve))
    except Exception as e:
        logger.error(f"Generate audio error: {str(e)}")
        return create_error_response(500, "Failed to generate audio.", str(e))

@api_view(['GET'])
def available_voices(request):
    logger.debug("Received available_voices request")
    try:
        voices = get_available_voices()
        if not voices:
            return create_error_response(500, "No voices available.")
        logger.debug(f"Returning {len(voices)} voice options")
        return JsonResponse({'voices': voices})
    except Exception as e:
        logger.error(f"Available voices error: {str(e)}")
        return create_error_response(500, "Failed to fetch available voices.", str(e))